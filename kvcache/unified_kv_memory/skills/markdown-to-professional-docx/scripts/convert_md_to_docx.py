#!/usr/bin/env python3
"""
Markdown to Professional DOCX Converter
Transforms technical design & requirement analysis Markdown documents into publication-grade Word files (.docx)
adhering to executive corporate design standards.
"""

import sys
import os
import re
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# --- COLOR PALETTE DEFINITIONS ---
NAVY_PRIMARY = (23, 54, 93)      # #17365D (Dark Executive Navy)
STEEL_SECONDARY = (47, 85, 151)   # #2F5597 (Steel Blue)
TEXT_DARK = (31, 31, 31)         # #1F1F1F (Charcoal Off-Black)
TEXT_MUTED = (102, 102, 102)     # #666666 (Gray)
WHITE = (255, 255, 255)          # #FFFFFF

HEX_NAVY_PRIMARY = "17365D"
HEX_STEEL_SECONDARY = "2F5597"
HEX_BG_ICE_BLUE = "EEF4F9"
HEX_BG_ALT_ROW = "F7F9FB"
HEX_BG_CODE = "F2F4F7"
HEX_BORDER_ACCENT = "D9E2F3"
HEX_BORDER_GRID = "D9D9D9"

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def add_p_border_bottom(p, color=HEX_BORDER_ACCENT, sz="5", space="3"):
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="{space}" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def add_p_left_border_and_bg(p, border_color=HEX_NAVY_PRIMARY, bg_color=HEX_BG_ICE_BLUE):
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="12" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
    pPr.append(shd)

def set_run_font(run, font_ascii="Arial", font_eastasia="SimSun", size_pt=10.5, color_rgb=TEXT_DARK, bold=False, italic=False):
    run.font.name = font_ascii
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_ascii}" w:hAnsi="{font_ascii}" w:eastAsia="{font_eastasia}"/>')
    rPr.append(rFonts)
    if size_pt:
        run.font.size = Pt(size_pt)
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)
    run.bold = bold
    run.italic = italic

def add_formatted_text(paragraph, text, default_font_ascii="Arial", default_font_eastasia="SimSun", default_size=10.5, default_color=TEXT_DARK):
    """Parses bold markdown `**text**` and adds runs."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            run = paragraph.add_run(content)
            set_run_font(run, font_ascii=default_font_ascii, font_eastasia=default_font_eastasia, size_pt=default_size, color_rgb=default_color, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, font_ascii=default_font_ascii, font_eastasia=default_font_eastasia, size_pt=default_size, color_rgb=default_color, bold=False)

def set_table_borders(table, color=HEX_BORDER_GRID, sz="4"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_page_setup(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

def setup_header_footer(doc, header_text=""):
    # Section 0 (Cover Page): no header/footer
    section0 = doc.sections[0]
    set_page_setup(section0)
    section0.different_first_page_header_footer = True

    # Main section
    section_body = doc.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
    set_page_setup(section_body)
    section_body.different_first_page_header_footer = False

    # Header
    header = section_body.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run(header_text)
    set_run_font(hrun, font_ascii="Arial", font_eastasia="SimSun", size_pt=9.0, color_rgb=TEXT_MUTED)

    # Footer
    footer = section_body.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run("—  页码  —")
    set_run_font(frun, font_ascii="Arial", font_eastasia="SimSun", size_pt=9.0, color_rgb=TEXT_MUTED)

    return section_body

def convert_md_to_docx(md_filepath, docx_filepath):
    with open(md_filepath, 'r', encoding='utf-8') as f:
        md_text = f.read()

    lines = md_text.splitlines()

    doc = docx.Document()

    # Document Defaults & Page setup
    doc.sections[0].page_width = Cm(21.0)
    doc.sections[0].page_height = Cm(29.7)
    doc.sections[0].top_margin = Cm(2.5)
    doc.sections[0].bottom_margin = Cm(2.3)
    doc.sections[0].left_margin = Cm(2.54)
    doc.sections[0].right_margin = Cm(2.54)
    doc.sections[0].different_first_page_header_footer = True

    title_text = ""
    blockquote_lines = []
    body_start_idx = 0

    # Extract Cover Page elements
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('# '):
            title_text = line[2:].strip()
            i += 1
            break
        i += 1

    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('>'):
            blockquote_lines.append(re.sub(r'^>\s*', '', line))
            i += 1
        elif not line:
            i += 1
        else:
            body_start_idx = i
            break

    # Build Cover Page
    category_p = doc.add_paragraph()
    category_p.paragraph_format.space_before = Pt(0)
    category_p.paragraph_format.space_after = Pt(20)
    c_run = category_p.add_run("技术架构与需求评审文件")
    set_run_font(c_run, font_ascii="Arial", font_eastasia="Microsoft YaHei", size_pt=11.0, color_rgb=STEEL_SECONDARY, bold=True)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(60)
    title_p.paragraph_format.space_after = Pt(12)
    t_run = title_p.add_run(title_text or "技术分析与需求规格说明")
    set_run_font(t_run, font_ascii="Arial", font_eastasia="Microsoft YaHei", size_pt=26.0, color_rgb=NAVY_PRIMARY, bold=True)

    if blockquote_lines:
        meta_table = doc.add_table(rows=0, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_table.autofit = False
        set_table_borders(meta_table, color="B0C4DE", sz="4")

        for b_line in blockquote_lines:
            parts = b_line.split('：', 1) if '：' in b_line else b_line.split(':', 1)
            row = meta_table.add_row()
            c0, c1 = row.cells[0], row.cells[1]
            c0.width = Cm(4.0)
            c1.width = Cm(12.0)
            set_cell_background(c0, HEX_BG_ICE_BLUE)
            set_cell_background(c1, "FFFFFF")
            set_cell_margins(c0, top=100, bottom=100, left=140, right=140)
            set_cell_margins(c1, top=100, bottom=100, left=140, right=140)

            p0 = c0.paragraphs[0]
            p0.paragraph_format.space_after = Pt(0)
            r0 = p0.add_run(parts[0].strip())
            set_run_font(r0, font_ascii="Arial", font_eastasia="Microsoft YaHei", size_pt=10.0, color_rgb=NAVY_PRIMARY, bold=True)

            p1 = c1.paragraphs[0]
            p1.paragraph_format.space_after = Pt(0)
            val_text = parts[1].strip() if len(parts) > 1 else ""
            r1 = p1.add_run(val_text)
            set_run_font(r1, font_ascii="Arial", font_eastasia="SimSun", size_pt=10.0, color_rgb=TEXT_DARK)

        p_spacer = doc.add_paragraph()
        p_spacer.paragraph_format.space_before = Pt(24)

    # Main Body Section
    body_section = setup_header_footer(doc, header_text=title_text)

    # Process remaining lines
    idx = body_start_idx
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    def flush_table(lines_list):
        if not lines_list:
            return
        rows_data = []
        for l in lines_list:
            cells = [c.strip() for c in l.strip('|').split('|')]
            # ignore separator row like |---|---|
            if all(re.match(r'^:?-+:?$', cell) for cell in cells if cell):
                continue
            rows_data.append(cells)

        if not rows_data:
            return

        cols_cnt = max(len(r) for r in rows_data)
        tbl = doc.add_table(rows=0, cols=cols_cnt)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True
        set_table_borders(tbl, color=HEX_BORDER_GRID, sz="4")

        for r_i, r_cells in enumerate(rows_data):
            row = tbl.add_row()
            # row split prevention & tblHeader
            trPr = row._tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
            if r_i == 0:
                trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

            is_header = (r_i == 0)
            is_even = (r_i % 2 == 0)

            for c_i in range(cols_cnt):
                cell = row.cells[c_i]
                set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
                cell_text = r_cells[c_i] if c_i < len(r_cells) else ""

                if is_header:
                    set_cell_background(cell, HEX_NAVY_PRIMARY)
                elif is_even:
                    set_cell_background(cell, HEX_BG_ALT_ROW)
                else:
                    set_cell_background(cell, "FFFFFF")

                cp = cell.paragraphs[0]
                cp.paragraph_format.space_before = Pt(2)
                cp.paragraph_format.space_after = Pt(2)

                if is_header:
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_formatted_text(cp, cell_text, default_font_ascii="Arial", default_font_eastasia="Microsoft YaHei", default_size=10.0, default_color=WHITE)
                    if cp.runs:
                        cp.runs[0].bold = True
                else:
                    add_formatted_text(cp, cell_text, default_font_ascii="Arial", default_font_eastasia="SimSun", default_size=9.5, default_color=TEXT_DARK)

        p_after = doc.add_paragraph()
        p_after.paragraph_format.space_before = Pt(4)
        p_after.paragraph_format.space_after = Pt(6)

    while idx < len(lines):
        line = lines[idx]
        s_line = line.strip()

        # Code block handling
        if s_line.startswith('```'):
            if in_code_block:
                # end block
                code_p = doc.add_paragraph()
                code_p.paragraph_format.space_before = Pt(4)
                code_p.paragraph_format.space_after = Pt(8)
                code_p.paragraph_format.line_spacing = 1.15
                add_p_left_border_and_bg(code_p, border_color="B0C4DE", bg_color=HEX_BG_CODE)

                code_content = "\n".join(code_lines)
                c_run = code_p.add_run(code_content)
                set_run_font(c_run, font_ascii="Consolas", font_eastasia="Microsoft YaHei", size_pt=8.5, color_rgb=TEXT_DARK)

                code_lines = []
                in_code_block = False
            else:
                if in_table:
                    flush_table(table_lines)
                    table_lines = []
                    in_table = False
                in_code_block = True
            idx += 1
            continue

        if in_code_block:
            code_lines.append(line)
            idx += 1
            continue

        # Table handling
        if s_line.startswith('|') and s_line.endswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(s_line)
            idx += 1
            continue
        else:
            if in_table:
                flush_table(table_lines)
                table_lines = []
                in_table = False

        # Blank line
        if not s_line:
            idx += 1
            continue

        # Headings
        if s_line.startswith('#'):
            h_match = re.match(r'^(#{1,4})\s+(.*)$', s_line)
            if h_match:
                level = len(h_match.group(1))
                h_text = h_match.group(2).strip()

                hp = doc.add_paragraph()

                if level == 1:
                    hp.paragraph_format.space_before = Pt(18)
                    hp.paragraph_format.space_after = Pt(8)
                    hp.paragraph_format.line_spacing = 1.15
                    hrun = hp.add_run(h_text)
                    set_run_font(hrun, font_ascii="Arial", font_eastasia="Microsoft YaHei", size_pt=16.0, color_rgb=NAVY_PRIMARY, bold=True)
                    add_p_border_bottom(hp, color=HEX_BORDER_ACCENT, sz="6")

                elif level == 2:
                    hp.paragraph_format.space_before = Pt(14)
                    hp.paragraph_format.space_after = Pt(6)
                    hp.paragraph_format.line_spacing = 1.15
                    hrun = hp.add_run(h_text)
                    set_run_font(hrun, font_ascii="Arial", font_eastasia="Microsoft YaHei", size_pt=14.0, color_rgb=NAVY_PRIMARY, bold=True)

                elif level == 3:
                    hp.paragraph_format.space_before = Pt(11)
                    hp.paragraph_format.space_after = Pt(4)
                    hp.paragraph_format.line_spacing = 1.15
                    hrun = hp.add_run(h_text)
                    set_run_font(hrun, font_ascii="Arial", font_eastasia="Microsoft YaHei", size_pt=12.0, color_rgb=STEEL_SECONDARY, bold=True)

                elif level == 4:
                    hp.paragraph_format.space_before = Pt(8)
                    hp.paragraph_format.space_after = Pt(4)
                    hp.paragraph_format.line_spacing = 1.15
                    hrun = hp.add_run(h_text)
                    set_run_font(hrun, font_ascii="Arial", font_eastasia="Microsoft YaHei", size_pt=11.0, color_rgb=STEEL_SECONDARY, bold=True)

                idx += 1
                continue

        # Blockquote Callout
        if s_line.startswith('>'):
            q_text = re.sub(r'^>\s*', '', s_line)
            qp = doc.add_paragraph()
            qp.paragraph_format.space_before = Pt(4)
            qp.paragraph_format.space_after = Pt(6)
            qp.paragraph_format.line_spacing = 1.3
            add_p_left_border_and_bg(qp, border_color=HEX_NAVY_PRIMARY, bg_color=HEX_BG_ICE_BLUE)
            add_formatted_text(qp, q_text, default_font_ascii="Arial", default_font_eastasia="Microsoft YaHei", default_size=10.0, default_color=NAVY_PRIMARY)
            idx += 1
            continue

        # Bullet List Items (- or *)
        if s_line.startswith('- ') or s_line.startswith('* '):
            item_text = s_line[2:].strip()
            bp = doc.add_paragraph()
            bp.paragraph_format.space_before = Pt(2)
            bp.paragraph_format.space_after = Pt(4)
            bp.paragraph_format.left_indent = Cm(0.75)
            brun = bp.add_run("•  ")
            set_run_font(brun, font_ascii="Arial", font_eastasia="Arial", size_pt=10.0, color_rgb=STEEL_SECONDARY, bold=True)
            add_formatted_text(bp, item_text, default_font_ascii="Arial", default_font_eastasia="SimSun", default_size=10.5, default_color=TEXT_DARK)
            idx += 1
            continue

        # Numbered List Items (1. 2. etc)
        num_match = re.match(r'^(\d+\.)\s+(.*)$', s_line)
        if num_match:
            prefix = num_match.group(1)
            item_text = num_match.group(2)
            np = doc.add_paragraph()
            np.paragraph_format.space_before = Pt(2)
            np.paragraph_format.space_after = Pt(4)
            np.paragraph_format.left_indent = Cm(0.75)
            nrun = np.add_run(prefix + " ")
            set_run_font(nrun, font_ascii="Arial", font_eastasia="Arial", size_pt=10.5, color_rgb=NAVY_PRIMARY, bold=True)
            add_formatted_text(np, item_text, default_font_ascii="Arial", default_font_eastasia="SimSun", default_size=10.5, default_color=TEXT_DARK)
            idx += 1
            continue

        # Standard Paragraph Body
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.45
        add_formatted_text(p, s_line, default_font_ascii="Arial", default_font_eastasia="SimSun", default_size=10.5, default_color=TEXT_DARK)
        idx += 1

    if in_table:
        flush_table(table_lines)

    doc.save(docx_filepath)
    print(f"Successfully converted '{md_filepath}' to professional DOCX at '{docx_filepath}'.")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python convert_md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)

    input_md = sys.argv[1]
    if len(sys.argv) >= 3:
        output_docx = sys.argv[2]
    else:
        output_docx = os.path.splitext(input_md)[0] + ".docx"

    convert_md_to_docx(input_md, output_docx)
